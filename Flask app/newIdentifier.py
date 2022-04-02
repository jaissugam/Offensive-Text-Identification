import numpy as np
import pickle
import re
import docx

def checkMessage(user_input):
    s = ' '.join(user_input)
    inputUser = re.split("[.?!]", s.strip())
    while "" in inputUser:
        inputUser.remove("")

    output = []
    for i in range(len(inputUser)):
        mes = [inputUser[i]]
        mes = np.array(mes)
        tfmodel=pickle.load(open('tfmodel.pkl','rb'))
        Tfidf_user = tfmodel.transform(mes)
        model=pickle.load(open('model.pkl','rb'))
        output.append(model.predict(Tfidf_user))
    return inputUser, output


def checkDocument(document):
    
    doc = docx.Document(document)
    result = [p.text for p in doc.paragraphs]
    result = result[0]
    output = []
    resultSentence = re.split('[?.!]', result.strip())
    for ele in resultSentence:
        if len(ele) < 1:
            resultSentence.remove(ele)
    tfmodel=pickle.load(open('tfmodel.pkl','rb'))
    Tfidf_user = tfmodel.transform(mes)
    model=pickle.load(open('model.pkl','rb'))
    output.append(model.predict(Tfidf_user))
    return resultSentence, output[0]


if __name__ == "__main__":
    userInput = [input("Enter Message: ").lower()]
    message, result = checkMessage(userInput)
    for i in range(len(message)):
        if result[i] == 0:
            print(message[i].strip(), ': NOT OFFENSIVE')
        else:
            print(message[i].strip(), ': OFFENSIVE')