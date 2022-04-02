import numpy as np
import pandas as pd
import timeit
import docx
import csv
import re
import string
import pickle
from pattern.text.en import lemma
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn import model_selection, naive_bayes, svm
from sklearn.metrics import accuracy_score
from sklearn.preprocessing import LabelEncoder
labels = []


def TextPreprocessing():
    # labels = []
    Data = []
    with open('olid-training-v1.tsv', 'r', encoding='utf-8') as text:
        reader = csv.reader(text, delimiter='\t')
        global labels
        labels = reader.__next__()
        for value in reader:
            Data.append(value)
    return Data


def Stop(Data):
    # removing stop words
    NoStopData = []
    stop = stopwords.words('english')
    for i in range(0, len(Data)):
        temp = []
        for word in Data[i][1].split(" "):
            if word not in stop:
                temp.append(word.lower())

        NoStopData.append(temp)
    return NoStopData


def Emojis(NoStopData):
    # removing emojis
    NoEmoStopData = []
    pattern = re.compile("["
                         u"\U0001F600-\U0001F64F"
                         u"\U0001F300-\U0001F5FF"
                         u"\U0001F680-\U0001F6FF"
                         u"\U0001F1E0-\U0001F1FF"
                         u"\U00002702-\U000027B0"
                         u"\U000024C2-\U0001F251"
                         u"\U0001f926-\U0001f937"
                         u"\U00010000-\U0010ffff"
                         u"\u2640-\u2642"
                         u"\u2600-\u2B55"
                         u"\u200d"
                         u"\u23cf"
                         u"\u23e9"
                         u"\u231a"
                         u"\ufe0f"
                         u"\u3030"
                         "]+", flags=re.UNICODE)

    for i in range(0, len(NoStopData)):
        temp = []
        for word in NoStopData[i]:
            temp.append(pattern.sub(r'', word))
        NoEmoStopData.append(temp)
    return NoEmoStopData


def ETC(NoEmoStopData):
    # removing punctuations, hashtags and @
    NoPuncEmoStopData = []
    prefix = ['@', '#']
    unnecessary = ['amp', 'url']
    for i in range(0, len(NoEmoStopData)):
        temp = []
        for word in NoEmoStopData[i]:
            for sep in string.punctuation:
                if sep not in prefix:
                    word = word.replace(sep, ' ')  # removing all punctuations other than @ or #

            if word != '':
                words = word.split()
                for w in words:
                    if w[0] not in prefix:  # removing words which have @ or # as prefix
                        if w not in unnecessary:  # removing the words url and amp which are unnecesary
                            temp.append(w.strip())

        NoPuncEmoStopData.append(temp)
    return NoPuncEmoStopData


def Digits(NoPuncEmoStopData):
    # removing digits
    NoDigitPuncEmoStopData = []
    for i in range(len(NoPuncEmoStopData)):
        temp = []
        for word in NoPuncEmoStopData[i]:
            if not word.isdigit():
                temp.append(word)
        NoDigitPuncEmoStopData.append(temp)
    return NoDigitPuncEmoStopData


def Lemmatization(Data, NoDigitPuncEmoStopData):
    # lemmatization
    FinalData = []
    i = 0

    for sentence in NoDigitPuncEmoStopData:
        newSentence = " ".join(sentence)
        lem = " "
        try:
            lem = " ".join([lemma(word) for word in newSentence.split()])
        except StopIteration:
            print("Error Happened")
        FinalData.append([lem, Data[i][2]])
        i += 1
    return FinalData
    # return labels, FinalData


def TrainModel():
    global labels
    Data = TextPreprocessing()
    NoStopData = Stop(Data)
    NoEmoStopData = Emojis(NoStopData)
    NoPuncEmoStopData = ETC(NoEmoStopData)
    NoDigitPuncEmoStopData = Digits(NoPuncEmoStopData)
    data = Lemmatization(Data, NoDigitPuncEmoStopData)

    labels = labels[1:3]
    x = []
    y = []
    for i in range(0, len(data)):
        x.append(data[i][0])
        y.append(data[i][1])

    Train_X, Test_X, Train_Y, Test_Y = model_selection.train_test_split(x, y, test_size=0.25)

    Encoder = LabelEncoder()
    Train_Y = Encoder.fit_transform(Train_Y)
    Test_Y = Encoder.fit_transform(Test_Y)

    Tf_idfModel = TfidfVectorizer()
    Tf_idfModel.fit(x)
    pickle.dump(Tf_idfModel,open('tfmodel.pkl','wb'))
    Tfidf_Train_X = Tf_idfModel.transform(Train_X)
    Tfidf_Test_X = Tf_idfModel.transform(Test_X)
    return Tfidf_Train_X, Train_Y, Tfidf_Test_X, Test_Y, Tf_idfModel


def checkMessage(user_input):
    TfidfTrainX, TrainY, TfidfTestX, TestY, Tf_idf = TrainModel()
    s = ' '.join(user_input)
    inputUser = re.split("[.?!]", s.strip())
    while "" in inputUser:
        inputUser.remove("")

    output = []
    for i in range(len(inputUser)):
        # mes = np.array(inputUser[i])
        mes = [inputUser[i]]
        mes = np.array(mes)
        Tfidf_user = Tf_idf.transform(mes)
        NaiveModel = naive_bayes.MultinomialNB()
        NaiveModel.fit(TfidfTrainX, TrainY)
        pickle.dump(NaiveModel,open('model.pkl','wb'))
        model=pickle.load(open('model.pkl','rb'))
        output.append(model.predict(Tfidf_user))
    return inputUser, output


def checkDocument(document):
    
    doc = docx.Document(document)
    TfidfTrainX, TrainY, TfidfTestX, TestY, Tf_idf = TrainModel()
    result = [p.text for p in doc.paragraphs]
    result = result[0]
    output = []
    resultSentence = re.split('[?.!]', result.strip())
    for ele in resultSentence:
        if len(ele) < 1:
            resultSentence.remove(ele)
    Tfidf_user = Tf_idf.transform(resultSentence)
    NaiveModel = naive_bayes.MultinomialNB()
    NaiveModel.fit(TfidfTrainX, TrainY)
    pickle.dump(NaiveModel,open('model.pkl','wb'))
    model=pickle.load(open('model.pkl','rb'))
    output.append(model.predict(Tfidf_user))
    return resultSentence, output[0]


if __name__ == "__main__":
    userInput = [input("Enter Message: ").lower()]
    message, result = checkMessage(userInput)
    for i in range(len(message)):
        if result[i] == 0:
            print(message[i].strip(), ': NOT OFFENSIVE')
        else:
            print(message[i].strip(), ': OFFENSIVE')


